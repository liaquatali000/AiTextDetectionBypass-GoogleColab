# First, install required packages:
# !pip install gradio nltk python-docx PyPDF2 transformers torch sentencepiece

import gradio as gr
import nltk
import docx
import PyPDF2
import os
import tempfile
from typing import Tuple, Dict, Optional
from transformers import pipeline
import shutil

# Download required NLTK data
nltk.download('punkt', quiet=True)

class TextProcessor:
    @staticmethod
    def read_file(file_path: str) -> str:
        """Read content from different file types."""
        ext = os.path.splitext(file_path)[1].lower()
        
        try:
            if ext == '.txt':
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif ext == '.docx':
                doc = docx.Document(file_path)
                return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            elif ext == '.pdf':
                text = ""
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                return text
            else:
                raise ValueError(f"Unsupported file format: {ext}")
        except Exception as e:
            raise Exception(f"Error reading file: {str(e)}")

    @staticmethod
    def save_file(content: str, output_path: str):
        """Save content to file."""
        try:
            ext = os.path.splitext(output_path)[1].lower()
            
            if ext == '.txt':
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            elif ext == '.docx':
                doc = docx.Document()
                for paragraph in content.split('\n'):
                    doc.add_paragraph(paragraph)
                doc.save(output_path)
            else:
                # Default to txt if format not supported
                output_path = os.path.splitext(output_path)[0] + '.txt'
                with open(output_path, 'w', encoding='utf-8') as f:
                    f.write(content)
        except Exception as e:
            raise Exception(f"Error saving file: {str(e)}")

class Paraphraser:
    def __init__(self):
        self.text_processor = TextProcessor()
        print("Loading paraphrasing model...")
        self.model = pipeline(
            "text2text-generation",
            model="tuner007/pegasus_paraphrase",
            device=-1  # Use CPU
        )
        print("Model loaded successfully!")

    def split_into_chunks(self, text: str, use_nltk: bool) -> list:
        """Split text into manageable chunks."""
        if use_nltk:
            return nltk.sent_tokenize(text)
        else:
            chunks = [s.strip() + '.' for s in text.split('.') if s.strip()]
            return chunks if chunks else [text]

    def paraphrase_text(self, text: str) -> str:
        """Paraphrase the input text."""
        try:
            # Handle empty or very short text
            if not text or len(text) < 10:
                return text

            # Truncate text if too long
            max_length = 512
            if len(text) > max_length:
                text = text[:max_length]

            result = self.model(
                text,
                max_length=100,
                num_beams=4,
                temperature=1.5
            )[0]['generated_text']
            
            return result
        except Exception as e:
            print(f"Error in paraphrasing: {str(e)}")
            return text

    def adjust_style(self, text: str, purpose: int, readability: int, tone: str) -> str:
        """Adjust text style based on parameters."""
        # Simplified style adjustments
        words = text.split()
        
        if readability == 1:  # High School
            # Simplify vocabulary
            if len(words) > 15:
                words = words[:15]
        elif readability == 3:  # Doctorate
            # Add academic tone
            if not text.endswith('.'):
                text += '.'
            
        if tone == "MORE_HUMAN":
            # Add conversational elements
            text = f"In other words, {text}"
        elif tone == "MORE_READABLE":
            # Break into shorter sentences
            if len(text) > 50:
                midpoint = text.find('.', len(text)//2)
                if midpoint != -1:
                    text = text[:midpoint+1] + '\n' + text[midpoint+1:]
        
        return text

    def process_file(self, purpose: int, readability: int, file_path: str, 
                    email: str, use_nltk: bool, save_same_format: bool, 
                    tone: str) -> str:
        """Process the entire file with paraphrasing."""
        try:
            # Read input file
            content = self.text_processor.read_file(file_path)
            
            # Split into chunks
            chunks = self.split_into_chunks(content, use_nltk)
            
            # Process each chunk
            processed_chunks = []
            total_chunks = len(chunks)
            
            for i, chunk in enumerate(chunks, 1):
                if len(chunk.strip()) > 0:
                    # Paraphrase
                    paraphrased = self.paraphrase_text(chunk)
                    # Adjust style
                    styled = self.adjust_style(paraphrased, purpose, readability, tone)
                    processed_chunks.append(styled)
                    
                    # Print progress
                    print(f"Processed chunk {i}/{total_chunks}")
            
            # Combine processed chunks
            final_text = ' '.join(processed_chunks)
            
            # Save output
            output_path = self._generate_output_path(file_path, save_same_format)
            self.text_processor.save_file(final_text, output_path)
            
            return f"Paraphrasing completed successfully! Output saved to: {output_path}"
            
        except Exception as e:
            return f"Error during processing: {str(e)}"

    def _generate_output_path(self, input_path: str, save_same_format: bool) -> str:
        """Generate output file path."""
        dir_name = os.path.dirname(input_path)
        base_name = os.path.splitext(os.path.basename(input_path))[0]
        ext = '.txt' if not save_same_format else os.path.splitext(input_path)[1]
        return os.path.join(dir_name, f"{base_name}_paraphrased{ext}")

class AIDetector:
    def __init__(self):
        pass

    def scan_text(self, text: str) -> Tuple[float, Dict]:
        """Scan text for AI content detection."""
        try:
            # Implement basic detection metrics
            detectors = {
                'Length Analysis': self._length_analysis(text),
                'Pattern Analysis': self._pattern_analysis(text),
                'Repetition Analysis': self._repetition_analysis(text)
            }
            
            # Calculate overall AI percentage
            total_ai = sum(d['ai_percentage'] for d in detectors.values())
            avg_ai = total_ai / len(detectors)
            
            return avg_ai, detectors
            
        except Exception as e:
            print(f"Error in AI detection: {str(e)}")
            return 0.0, {}

    def _length_analysis(self, text: str) -> Dict:
        """Analyze sentence length distribution."""
        sentences = nltk.sent_tokenize(text)
        avg_length = sum(len(s.split()) for s in sentences) / len(sentences) if sentences else 0
        
        # Longer average sentence length might indicate AI
        ai_percentage = min(100, max(0, (avg_length - 10) * 5))
        return {
            'ai_percentage': ai_percentage,
            'human_percentage': 100 - ai_percentage
        }

    def _pattern_analysis(self, text: str) -> Dict:
        """Analyze text patterns."""
        words = text.lower().split()
        unique_words = len(set(words))
        total_words = len(words)
        
        # Higher variety might indicate AI
        variety_ratio = (unique_words / total_words) if total_words > 0 else 0
        ai_percentage = min(100, max(0, variety_ratio * 100))
        return {
            'ai_percentage': ai_percentage,
            'human_percentage': 100 - ai_percentage
        }

    def _repetition_analysis(self, text: str) -> Dict:
        """Analyze word repetition patterns."""
        words = text.lower().split()
        word_freq = {}
        for word in words:
            word_freq[word] = word_freq.get(word, 0) + 1
        
        # More even distribution might indicate AI
        max_freq = max(word_freq.values()) if word_freq else 0
        ai_percentage = min(100, max(0, (1 - (max_freq / len(words))) * 100)) if words else 50
        return {
            'ai_percentage': ai_percentage,
            'human_percentage': 100 - ai_percentage
        }

def create_interface():
    paraphraser = Paraphraser()
    ai_detector = AIDetector()

    def paraphrase_text(purpose, readability, tone, uploaded_file, email, use_nltk, save_same_format):
        if uploaded_file is None:
            return "Please upload a file"
        
        try:
            # Get the temporary file path from Gradio's file object
            temp_path = uploaded_file.name
            
            # Convert inputs
            purpose_choice = ["General Writing", "Essay", "Article", "Marketing Material", 
                            "Story", "Cover letter", "Report", "Business Material", 
                            "Legal Material"].index(purpose) + 1
            
            readability_choice = ["High School", "University", "Doctorate", 
                                "Journalist", "Marketing"].index(readability) + 1
            
            tone_choice = tone.upper().replace(' ', '_')
            
            # Process the file
            result = paraphraser.process_file(
                purpose_choice, readability_choice, temp_path, 
                email, use_nltk, save_same_format, tone_choice
            )
            
            return result
        except Exception as e:
            return f"An error occurred: {str(e)}"

    def scan_ai_content(text):
        if not text.strip():
            return "Please enter text to scan."
        
        try:
            detection_response = ai_detector.scan_text(text)
            if detection_response is None:
                return "Failed to detect AI in the text."
            
            ai_percentage, detector_results = detection_response
            
            # Format results
            result_text = f"Overall AI Content: {ai_percentage:.1f}%\n\nDetailed Results:\n"
            for detector, result in detector_results.items():
                ai_perc = result.get('ai_percentage', 0)
                human_perc = result.get('human_percentage', 100)
                result_text += f"\n{detector}:\n"
                result_text += f"AI: {ai_perc:.1f}% | Human: {human_perc:.1f}%"
            
            return result_text
        except Exception as e:
            return f"Error during scanning: {str(e)}"

    # Create the interface
    with gr.Blocks(title="ParaGenie V2.0", theme=gr.themes.Soft()) as app:
        gr.Markdown("""
        # ParaGenie V2.0
        ### Paraphrase and analyze text with AI
        """)
        
        with gr.Tabs():
            # Paraphraser Tab
            with gr.Tab("Paraphraser"):
                with gr.Column():
                    purpose = gr.Dropdown(
                        choices=["General Writing", "Essay", "Article", "Marketing Material", 
                                "Story", "Cover letter", "Report", "Business Material", 
                                "Legal Material"],
                        label="Purpose of Writing",
                        value="General Writing"
                    )
                    readability = gr.Dropdown(
                        choices=["High School", "University", "Doctorate", 
                                "Journalist", "Marketing"],
                        label="Readability Level",
                        value="University"
                    )
                    tone = gr.Dropdown(
                        choices=["Balanced", "More Human", "More Readable"],
                        label="Tone",
                        value="Balanced"
                    )
                    file_input = gr.File(
                        label="Upload Document",
                        file_types=[".txt", ".docx", ".pdf"]
                    )
                    email = gr.Textbox(
                        label="Email Address",
                        placeholder="Enter your email address"
                    )
                    with gr.Row():
                        use_nltk = gr.Checkbox(label="Use NLTK for splitting chunks")
                        save_format = gr.Checkbox(label="Save file in the same format")
                    
                    submit_btn = gr.Button("Start Paraphrasing")
                    output = gr.Textbox(label="Results")
                    
                    submit_btn.click(
                        fn=paraphrase_text,
                        inputs=[purpose, readability, tone, file_input, 
                               email, use_nltk, save_format],
                        outputs=output
                    )
            
            # AI Scanner Tab
            with gr.Tab("AI Scanner"):
                with gr.Column():
                    text_input = gr.Textbox(
                        label="Enter or paste text to scan",
                        lines=10,
                        placeholder="Enter or paste text here..."
                    )
                    scan_btn = gr.Button("Scan")
                    scan_output = gr.Textbox(label="Scan Results")
                    
                    scan_btn.click(
                        fn=scan_ai_content,
                        inputs=text_input,
                        outputs=scan_output
                    )
    
    return app

# For running in Colab
if __name__ == "__main__":
    app = create_interface()
    app.launch(share=True)